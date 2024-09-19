import docx
import os
from docx import document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx import document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from tkinter import filedialog
from tkinter import messagebox

from excel_data import *


def new_page(doc_in, page_index, in_data):

    if page_index != 1:
        doc_in.add_page_break()

    # 2.添加 表格
    table = doc_in.add_table(rows=25, cols=10)

    # 设置表格的线条为细线条
    table.style = 'Table Grid'

    # 合并单元格 0,0 到 0,2
    a = table.cell(0, 0).merge(table.cell(0, 2))
    a.text = '样品型号和规格'
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 合并单元格 0,3 到 0,4
    table.cell(0, 3).merge(table.cell(0, 5))

    # 合并单元格 0,6 到 0,7
    table.cell(0, 6).text = '报告编号'
    c = table.cell(0, 6).merge(table.cell(0, 7))
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.cell(0, 8).merge(table.cell(0, 9))

    # 第二行，合并单元格 1,1 到 1,2 和 1,4 到 1,5
    table.cell(1, 0).text = '序号'
    row2_1 = table.cell(1, 1).merge(table.cell(1, 3))
    row2_1.text = '检 测 项 目'

    table.cell(1, 4).text = '单位'
    table.cell(1, 5).text = '技 术 要 求'
    # 设置宽度
    table.cell(1, 5).width = Cm(2.5)
    row2_2 = table.cell(1, 6).merge(table.cell(1, 8))
    row2_2.text = '检 测 结 果'
    table.cell(1, 9).text = '单项评定'

    # 从第三行开始，合并单元格 2,1 到 2,3 和 2,6 到 2,8
    for i, _ in enumerate(table.rows[2:]):
        table.cell(i + 2, 1).merge(table.cell(i + 2, 3))
        table.cell(i + 2, 6).merge(table.cell(i + 2, 8))

    # 设置字体为宋体，小四
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'SimSun'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                    run.font.size = Pt(10.5)

    # 设置第一二行高为 1.42cm
    for row in table.rows[:2]:
        row.height = Cm(1.5)
    # 之后的行高为0.5cm
    for row in table.rows[2:]:
        row.height = Cm(0.8)

    # 设置字体垂直居中
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 从第三行开始写入数据
    # 能写入数据的列为 1, 4, 5, 6, 9
    write_col = [0, 1, 4, 5, 6, 9]

    # 表格填充数据
    for i, data in enumerate(in_data):
        for j, value in enumerate(data):
            print(f"正在第{page_index}页 第{i}行 第{j}列")
            table.cell(i + 2, write_col[j]).text = value

    return doc_in


def write_doc(save_path, in_data, element_log):
    # 1.新建 文档对象
    doc = docx.Document()
    assert isinstance(doc, document.Document)  # 格式：doc, doc 的类型

    # 每次写几行
    write_page_row = 23
    # 判断传几次
    page_num = len(in_data) // write_page_row
    for i in range(page_num):
        element_log.insert('end', f'正在写入第{i + 1}页-----共{page_num}页\n')
        # 刷新主窗口
        element_log.update()
        element_log.see('end')
        doc = new_page(doc, i + 1, in_data[i * write_page_row: (i + 1) * write_page_row])

    # 循环每一页，添加页眉
    page_all = len(doc.sections)
    for page_index, section in enumerate(doc.sections):
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = f"共 {page_all} 页  第 {page_index} 页"
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # 设置字体大小为 12
        for run in header_paragraph.runs:
            run.font.size = Pt(8.5)

        # 添加页脚
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = "注：“单项评定”符号含义：P：检测结果符合要求；F：检测结果不符合要求；N：检测结果不要求判定。"
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 设置字体大小为 8.5
        for run in footer_paragraph.runs:
            run.font.size = Pt(8.5)

    doc.save(save_path)


def open_folder(folder_path):
    folder_selected = filedialog.askopenfilename()
    folder_path.set("您选的数据文件为： " + folder_selected)
    return folder_selected


def convert_data(file_path_in, element_log):
    # 如果 file_path_in 为 请选择数据文件，tkinter弹窗提示
    if file_path_in.get() == "请选择数据文件":
        messagebox.showinfo("提示", "请选择数据文件后再点击开始", icon='warning')
        return
    file_path_in = file_path_in.get().split('： ')[1]

    # 清洗数据
    element_log.insert('end', '开始清洗数据\n')
    element_log.insert('end', '正在清洗数据\n')
    element_log.see('end')

    # 刷新主窗口
    element_log.update()

    res_in = clean_data(file_path_in)
    # 转为二维数组
    out_data = []
    for col in res_in:
        out_data.append([col.seq, col.item, col.unit, col.requirement, col.result, col.assessment])

    # 输出地址
    # put_path = './output.docx'
    put_path = os.path.dirname(file_path_in) + '/output.docx'

    # 写入数据到 docx
    element_log.insert('end', '开始写入数据\n')
    element_log.insert('end', '正在写入数据\n')
    # 刷新主窗口
    element_log.update()
    write_doc(put_path, out_data, element_log)

    element_log.insert('end', '数据写入完成\n保存地址为{} 请查看\n'.format(put_path))
    # 刷新主窗口
    element_log.update()
    element_log.see('end')

